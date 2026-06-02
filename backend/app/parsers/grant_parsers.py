from __future__ import annotations
from typing import Dict, Any, Tuple
from io import BytesIO
import json
import logging
import os
import re

logger = logging.getLogger(__name__)
CHAT_MODEL = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")
VALID_SECTION_ROMANS = {"I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"}
QUESTION_ID_PATTERN = r"(?:Q\d+[a-z]?|\d+(?:\.[A-Za-z0-9]+)*)"
PROMPT_MATCH_STOPWORDS = {
    "the", "and", "for", "that", "this", "with", "from", "your", "please", "provide",
    "enter", "share", "what", "which", "about", "more", "than", "into", "over",
    "have", "has", "will", "would", "should", "could", "their", "organization",
    "project", "application", "below", "following", "information", "select", "describe",
    "explain", "list", "identify", "required", "response", "question", "public", "services",
}

def _normalize_text_snippet(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _matching_tokens(text: str) -> list[str]:
    words = [w.lower() for w in re.findall(r"[A-Za-z][A-Za-z/&'-]{2,}", text or "")]
    return [word for word in words if word not in PROMPT_MATCH_STOPWORDS][:12]


def _read_uploaded_file_bytes(uploaded_file) -> bytes:
    if hasattr(uploaded_file, "getvalue"):
        try:
            data = uploaded_file.getvalue()
            return data if isinstance(data, (bytes, bytearray)) else bytes(data or b"")
        except Exception:
            pass
    if hasattr(uploaded_file, "read"):
        try:
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            data = uploaded_file.read()
            if hasattr(uploaded_file, "seek"):
                uploaded_file.seek(0)
            return data if isinstance(data, (bytes, bytearray)) else bytes(data or b"")
        except Exception:
            return b""
    return b""


def _read_txt(file) -> str:
    return file.getvalue().decode("utf-8", errors="ignore")

def _read_pdf(file) -> str:
    """
    Best-effort PDF text extraction.
    Tries PyPDF2/pypdf first, then pdfplumber (if installed) as a fallback.
    """
    parts: list[str] = []
    stream = file
    if hasattr(file, "getvalue"):
        try:
            stream = BytesIO(file.getvalue())
        except Exception:
            stream = file

    reader_cls = None
    try:
        import PyPDF2
        reader_cls = PyPDF2.PdfReader
    except Exception:
        try:
            from pypdf import PdfReader as PypdfReader
            reader_cls = PypdfReader
        except Exception:
            reader_cls = None

    if reader_cls is not None:
        try:
            reader = reader_cls(stream)
            for page in reader.pages:
                parts.append(page.extract_text() or "")
        except Exception:
            parts = []

    text = "\n".join(parts).strip()
    if len(text) > 200:
        return text

    # Fallback parser for PDFs where PyPDF2 misses ordering/blocks
    try:
        # rewind underlying buffer if available
        if hasattr(stream, "seek"):
            stream.seek(0)
        import pdfplumber

        parts = []
        with pdfplumber.open(stream) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()
    except Exception:
        return text

def _read_docx(file) -> str:
    try:
        import docx  # python-docx
    except Exception:
        return ""
    try:
        if hasattr(file, "seek"):
            file.seek(0)
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception:
        return ""


def extract_text_from_upload(uploaded_file) -> str:
    """Best-effort text extraction for supporting context documents."""
    name = str(getattr(uploaded_file, "name", "") or "").lower()
    if name.endswith(".pdf"):
        return _read_pdf(uploaded_file)
    if name.endswith(".docx"):
        return _read_docx(uploaded_file)
    if name.endswith((".txt", ".md", ".csv", ".json")):
        return _read_txt(uploaded_file)
    return _read_txt(uploaded_file)


def _document_ai_is_configured() -> bool:
    return bool(
        os.getenv("GOOGLE_CLOUD_PROJECT")
        and os.getenv("DOCUMENTAI_LOCATION")
        and os.getenv("DOCUMENTAI_PROCESSOR_ID")
    )


def _docai_text_from_anchor(full_text: str, text_anchor: Any) -> str:
    if not full_text or not text_anchor or not getattr(text_anchor, "text_segments", None):
        return ""
    parts: list[str] = []
    for segment in text_anchor.text_segments:
        try:
            start = int(getattr(segment, "start_index", 0) or 0)
            end = int(getattr(segment, "end_index", 0) or 0)
        except Exception:
            continue
        if end <= start:
            continue
        parts.append(full_text[start:end])
    return _normalize_text_snippet(" ".join(parts))


def _docai_layout_text(full_text: str, layout: Any) -> str:
    if not layout:
        return ""
    return _docai_text_from_anchor(full_text, getattr(layout, "text_anchor", None))


def _build_pdf_shards(file_bytes: bytes, pages_per_shard: int = 15) -> list[bytes]:
    try:
        from pypdf import PdfReader, PdfWriter
    except Exception:
        try:
            import PyPDF2
            PdfReader = PyPDF2.PdfReader
            PdfWriter = PyPDF2.PdfWriter
        except Exception:
            return [file_bytes]

    try:
        reader = PdfReader(BytesIO(file_bytes))
        total_pages = len(reader.pages)
        if total_pages <= pages_per_shard:
            return [file_bytes]

        shards: list[bytes] = []
        for start in range(0, total_pages, pages_per_shard):
            writer = PdfWriter()
            for page_index in range(start, min(start + pages_per_shard, total_pages)):
                writer.add_page(reader.pages[page_index])
            buffer = BytesIO()
            writer.write(buffer)
            shards.append(buffer.getvalue())
        return shards or [file_bytes]
    except Exception:
        return [file_bytes]


def _extract_document_ai_payload(
    client: Any,
    documentai: Any,
    processor_name: str,
    file_bytes: bytes,
    mime_type: str,
) -> tuple[str, list[str], int, int, int, list[dict[str, str]]]:
    raw_document = documentai.RawDocument(content=file_bytes, mime_type=mime_type)
    result = client.process_document(
        request=documentai.ProcessRequest(
            name=processor_name,
            raw_document=raw_document,
            imageless_mode=True,
        )
    )
    document = result.document
    full_text = str(getattr(document, "text", "") or "").strip()
    augmented_blocks: list[str] = [full_text] if full_text else []
    form_field_count = 0
    table_count = 0
    page_count = 0
    extracted_fields: list[dict[str, str]] = []

    for page in getattr(document, "pages", []) or []:
        page_count += 1
        for form_field in getattr(page, "form_fields", []) or []:
            key_text = _docai_layout_text(full_text, getattr(form_field, "field_name", None))
            value_text = _docai_layout_text(full_text, getattr(form_field, "field_value", None))
            if not key_text:
                continue
            form_field_count += 1
            if value_text:
                extracted_fields.append({"key_text": key_text, "value_text": value_text})
            if value_text:
                augmented_blocks.append(f"{key_text}\nRESPONSE: {value_text}")
            else:
                augmented_blocks.append(f"{key_text}\nRESPONSE:")

        for table in getattr(page, "tables", []) or []:
            table_count += 1
            row_lines: list[str] = []
            header_cells = []
            for header_row in getattr(table, "header_rows", []) or []:
                header_values = [
                    _docai_layout_text(full_text, getattr(cell, "layout", None))
                    for cell in getattr(header_row, "cells", []) or []
                ]
                header_cells.extend([value for value in header_values if value])
            if header_cells:
                row_lines.append(" | ".join(header_cells))
            for body_row in getattr(table, "body_rows", []) or []:
                body_values = [
                    _docai_layout_text(full_text, getattr(cell, "layout", None))
                    for cell in getattr(body_row, "cells", []) or []
                ]
                body_values = [value for value in body_values if value]
                if body_values:
                    row_lines.append(" | ".join(body_values))
            if row_lines:
                augmented_blocks.append("\n".join(row_lines))

    return full_text, augmented_blocks, page_count, form_field_count, table_count, extracted_fields


def _score_extracted_field_for_prompt(prompt_item: dict[str, Any], field: dict[str, str]) -> int:
    prompt_id = str(prompt_item.get("prompt_id") or "").strip().lower()
    prompt_text = str(prompt_item.get("prompt_text") or "")
    field_key = str(field.get("key_text") or "")
    field_value = str(field.get("value_text") or "").strip()
    if not field_key or not field_value:
        return 0

    prompt_tokens = _matching_tokens(prompt_text)
    field_tokens = set(_matching_tokens(field_key))
    score = sum(2 for token in prompt_tokens if token in field_tokens)

    prompt_text_norm = _normalize_text_snippet(prompt_text).lower()
    field_key_norm = _normalize_text_snippet(field_key).lower()

    if prompt_id and prompt_id in field_key_norm:
        score += 4
    if prompt_text_norm and prompt_text_norm[:50] in field_key_norm:
        score += 4
    if "email" in prompt_text_norm and "@" in field_value:
        score += 3
    if "website" in prompt_text_norm and re.search(r"https?://|www\.", field_value, flags=re.I):
        score += 3
    if "budget" in prompt_text_norm and re.search(r"\$\s?\d", field_value):
        score += 3
    if "year" in prompt_text_norm and re.search(r"\b(?:19|20)\d{2}\b", field_value):
        score += 2
    if "employee" in prompt_text_norm and re.search(r"\b\d+\b", field_value):
        score += 2
    return score


def _bind_extracted_values_to_sections(sections: list[dict], extracted_fields: list[dict[str, str]]) -> list[dict]:
    if not sections or not extracted_fields:
        return sections

    for section in sections:
        prompt_items = section.get("prompt_items") or []
        bound_prompt_items: list[dict[str, Any]] = []
        for prompt_item in prompt_items:
            updated = dict(prompt_item)
            if updated.get("response_value"):
                bound_prompt_items.append(updated)
                continue
            ranked = sorted(
                extracted_fields,
                key=lambda field: _score_extracted_field_for_prompt(updated, field),
                reverse=True,
            )
            best = ranked[0] if ranked else None
            best_score = _score_extracted_field_for_prompt(updated, best) if best else 0
            if best and best_score >= 4:
                updated["response_value"] = str(best.get("value_text") or "").strip()
                updated["source_origin"] = "document_ai_form_field"
                updated["source_confidence"] = "high" if best_score >= 7 else "medium"
            bound_prompt_items.append(updated)
        section["prompt_items"] = bound_prompt_items
    return sections


def _extract_sections_with_document_ai(file_bytes: bytes, mime_type: str) -> tuple[list[dict], str, dict[str, Any], str | None]:
    project_id = os.getenv("GOOGLE_CLOUD_PROJECT")
    location = os.getenv("DOCUMENTAI_LOCATION")
    processor_id = os.getenv("DOCUMENTAI_PROCESSOR_ID")
    if not project_id or not location or not processor_id:
        return [], "", {}, "document_ai_not_configured"
    if not file_bytes:
        return [], "", {}, "empty_file_bytes"

    try:
        from google.api_core.client_options import ClientOptions
        from google.cloud import documentai
    except Exception as exc:
        return [], "", {}, f"documentai_sdk_unavailable:{type(exc).__name__}"

    try:
        endpoint = f"{location}-documentai.googleapis.com"
        client = documentai.DocumentProcessorServiceClient(
            client_options=ClientOptions(api_endpoint=endpoint)
        )
        processor_name = client.processor_path(project_id, location, processor_id)
    except Exception as exc:
        logger.exception("Document AI processing failed")
        return [], "", {}, f"{type(exc).__name__}: {exc}"

    shard_count = 1
    total_pages = 0
    total_form_fields = 0
    total_tables = 0
    combined_blocks: list[str] = []
    full_text_parts: list[str] = []
    extracted_fields: list[dict[str, str]] = []

    try:
        shards = _build_pdf_shards(file_bytes, pages_per_shard=15) if mime_type == "application/pdf" else [file_bytes]
        shard_count = len(shards)
        for shard_bytes in shards:
            full_text, augmented_blocks, page_count, form_field_count, table_count, shard_fields = _extract_document_ai_payload(
                client,
                documentai,
                processor_name,
                shard_bytes,
                mime_type,
            )
            if full_text:
                full_text_parts.append(full_text)
            combined_blocks.extend(augmented_blocks)
            extracted_fields.extend(shard_fields)
            total_pages += page_count
            total_form_fields += form_field_count
            total_tables += table_count
    except Exception as exc:
        logger.exception("Document AI processing failed")
        return [], "", {}, f"{type(exc).__name__}: {exc}"

    full_text = "\n\n".join(part for part in full_text_parts if part).strip()
    augmented_text = "\n\n".join(block for block in combined_blocks if block).strip()
    structure_text = full_text or augmented_text
    sections = _extract_sections_from_text(structure_text)
    sections = _bind_extracted_values_to_sections(sections, extracted_fields)
    meta = {
        "page_count": total_pages,
        "form_field_count": total_form_fields,
        "table_count": total_tables,
        "extracted_field_count": len(extracted_fields),
        "shard_count": shard_count,
        "used_augmented_text": bool(structure_text == augmented_text and augmented_text and augmented_text != full_text),
        "processor_id": processor_id,
        "location": location,
    }
    return sections, structure_text, meta, None


def _is_valid_roman_heading(line: str) -> bool:
    match = re.match(r"^\s*([IVXLCDM]{1,8})[\)\.\-:]\s+\S+", line or "", flags=re.I)
    if not match:
        return False
    return match.group(1).upper() in VALID_SECTION_ROMANS

def _is_probable_heading(line: str) -> bool:
    ln = line.strip()
    if not ln:
        return False
    if len(ln) < 3 or len(ln) > 140:
        return False
    low = ln.lower()

    # Exclude obvious body lines / bullets.
    if ln.startswith(("-", "*", "\u2022")):
        return False
    if len(ln.split()) > 18:
        return False
    if re.match(r"^page\s+\d+\s*(?:of|/)\s*\d+\b", low):
        return False
    if re.search(r"\b(?:https?://|www\.)\S+", ln):
        return False
    if re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", ln, flags=re.I):
        return False
    if re.match(r"^[\(\[]?\s*[xX ]\s*[\)\]]\s+\S+", ln):
        return False

    # Strong heading patterns with explicit numbering systems.
    if _is_valid_roman_heading(ln):
        return True
    if re.match(r"^\d+(\.\d+)*[\)\.\-:]?\s+\S+", ln):
        return True
    if re.match(r"^[A-Z][\)\.\-:]\s+\S+", ln):  # A) Scope
        return True
    if re.match(r"^(SECTION|PART|APPENDIX)\s+[A-Z0-9]+", ln, flags=re.I):
        return True
    if ln.endswith(":") and len(ln.split()) <= 12:
        return True

    # All-caps headings are common in grant docs.
    letters = [c for c in ln if c.isalpha()]
    if letters:
        upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
        if upper_ratio >= 0.8 and len(ln.split()) <= 10 and not re.search(r"[,.]", ln):
            return True

    # Title Case headings are much noisier in grant PDFs, so keep this narrow.
    words = re.findall(r"[A-Za-z][A-Za-z/&'-]*", ln)
    if 2 <= len(words) <= 4:
        title_like = sum(1 for w in words if w[0].isupper())
        if title_like / max(1, len(words)) >= 0.95 and not re.search(r"[.!?,;]$", ln):
            return True

    return False

def _is_top_level_section_heading(line: str) -> bool:
    ln = line.strip()
    if not ln:
        return False
    if _is_valid_roman_heading(ln):
        return True
    if re.match(r"^\d+(\.\d+){0,2}[\)\.\-:]?\s+\S+", ln):
        return True
    if re.match(r"^(SECTION|PART|APPENDIX)\s+[A-Z0-9]+", ln, flags=re.I):
        return True
    return False


def _is_explicit_non_numeric_top_level_heading(line: str) -> bool:
    ln = line.strip()
    if not ln:
        return False
    if _is_valid_roman_heading(ln):
        return True
    if re.match(r"^(SECTION|PART|APPENDIX)\s+[A-Z0-9]+", ln, flags=re.I):
        return True
    return False


def _parse_single_level_numbered_heading(line: str) -> tuple[int, str] | None:
    match = re.match(r"^\s*(\d+)[\)\.\-:]\s+(.+?)\s*$", line or "")
    if not match:
        return None
    try:
        number = int(match.group(1))
    except Exception:
        return None
    title = match.group(2).strip()
    if not title:
        return None
    return number, title


def _base_numeric_heading_score(lines: list[str], index: int) -> float:
    line = lines[index].strip()
    parsed = _parse_single_level_numbered_heading(line)
    if not parsed:
        return float("-inf")

    _, title = parsed
    title_words = len(title.split())
    score = 0.0

    if 1 <= title_words <= 8:
        score += 2.5
    elif title_words <= 12:
        score += 1.0
    else:
        score -= 2.5

    if title.endswith(":"):
        score -= 1.0

    prev_line = lines[index - 1].strip() if index > 0 else ""
    prev_non_empty = ""
    for cursor in range(index - 1, -1, -1):
        if lines[cursor].strip():
            prev_non_empty = lines[cursor].strip()
            break

    next_non_empty = ""
    for cursor in range(index + 1, len(lines)):
        if lines[cursor].strip():
            next_non_empty = lines[cursor].strip()
            break

    if not prev_line:
        score += 1.5
    if prev_non_empty and prev_non_empty.endswith((".", ":", ";")):
        score += 0.5

    if next_non_empty:
        if next_non_empty.startswith(("o ", "•", "-", "*")):
            score -= 2.0
        elif _parse_single_level_numbered_heading(next_non_empty):
            score -= 3.0
        else:
            score += 1.5

    if re.search(r"\b(?:goal|objectives?|details?|overview|challenge|journey|program|roles?|budget|cost|timeline|evaluation|plan)\b", title, flags=re.I):
        score += 1.5

    return score


def _select_primary_numbered_heading_indices(lines: list[str]) -> list[int]:
    candidates: list[dict[str, float | int]] = []
    for idx, line in enumerate(lines):
        parsed = _parse_single_level_numbered_heading(line.strip())
        if not parsed:
            continue
        number, _ = parsed
        candidates.append(
            {
                "index": idx,
                "number": number,
                "score": _base_numeric_heading_score(lines, idx),
            }
        )

    if len(candidates) < 3:
        return []

    best_score = [float(item["score"]) for item in candidates]
    best_length = [1 for _ in candidates]
    previous = [-1 for _ in candidates]

    for i in range(len(candidates)):
        current_number = int(candidates[i]["number"])
        current_index = int(candidates[i]["index"])
        current_base = float(candidates[i]["score"])
        for j in range(i):
            prior_number = int(candidates[j]["number"])
            prior_index = int(candidates[j]["index"])
            if prior_index >= current_index:
                continue
            if prior_number + 1 != current_number:
                continue

            gap_bonus = min(2.0, max(0.0, (current_index - prior_index) / 80.0))
            candidate_score = best_score[j] + current_base + gap_bonus
            candidate_length = best_length[j] + 1
            if (
                candidate_length > best_length[i]
                or (candidate_length == best_length[i] and candidate_score > best_score[i])
            ):
                best_score[i] = candidate_score
                best_length[i] = candidate_length
                previous[i] = j

    best_idx = max(range(len(candidates)), key=lambda i: (best_length[i], best_score[i]))
    if best_length[best_idx] < 3:
        return []

    chain: list[int] = []
    cursor = best_idx
    while cursor != -1:
        chain.append(int(candidates[cursor]["index"]))
        cursor = previous[cursor]
    return list(reversed(chain))

def _is_heading_continuation(line: str) -> bool:
    ln = line.strip()
    if not ln:
        return False
    if _is_top_level_section_heading(ln):
        return False
    if len(ln) > 60:
        return False
    if re.search(r"\d", ln):
        return False
    words = re.findall(r"[A-Za-z&/\-']+", ln)
    if not words or len(words) > 4:
        return False
    if words[0].upper() not in {"OF", "AND", "FOR", "TO", "IN", "ON"}:
        return False
    letters = [c for c in ln if c.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return upper_ratio >= 0.8

def _extract_sections_from_numbered_headings(lines: list[str]) -> list[dict]:
    sections = []
    seen_keys: dict[str, int] = {}
    heading_idxs: list[int] = []
    headings: list[str] = []

    primary_numeric_idxs = set(_select_primary_numbered_heading_indices(lines))

    i = 0
    while i < len(lines):
        ln = lines[i].strip()
        is_primary_numeric = i in primary_numeric_idxs
        is_other_top_level = _is_explicit_non_numeric_top_level_heading(ln)
        if is_primary_numeric or is_other_top_level:
            heading = ln.rstrip(":").strip()
            j = i + 1
            continuation_parts: list[str] = []
            while j < len(lines) and _is_heading_continuation(lines[j]):
                continuation_parts.append(lines[j].strip().rstrip(":"))
                j += 1
            if continuation_parts:
                heading = f"{heading} {' '.join(continuation_parts)}".strip()
            heading_idxs.append(i)
            headings.append(heading)
            i = j
            continue
        i += 1

    if len(heading_idxs) < 2:
        return []

    for idx, start in enumerate(heading_idxs):
        end = heading_idxs[idx + 1] if idx + 1 < len(heading_idxs) else len(lines)
        heading = headings[idx]
        body_lines = []
        cursor = start + 1
        while cursor < end and _is_heading_continuation(lines[cursor]):
            cursor += 1

        for raw_line in lines[cursor:end]:
            body_line = raw_line.strip()
            if not body_line:
                continue
            if body_line.upper() == "DISCLAIMER":
                break
            if _is_probable_heading(body_line) and not body_lines:
                # Skip table labels or cover metadata immediately beneath a true section heading.
                continue
            body_lines.append(body_line)

        body = "\n".join(body_lines).strip()
        key_base = re.sub(r"[^a-z0-9]+", "_", heading.lower()).strip("_")[:60] or f"section_{idx+1}"
        count = seen_keys.get(key_base, 0) + 1
        seen_keys[key_base] = count
        key = key_base if count == 1 else f"{key_base}_{count}"

        sec = {"key": key, "title": heading, "guidance": body[:3000]}
        wl = _extract_word_limit(f"{heading}\n{body}")
        if wl:
            sec["word_limit"] = wl
        sections.append(sec)

    return _normalize_sections(sections)


def _is_question_heading(line: str) -> bool:
    ln = (line or "").strip()
    if not ln:
        return False
    if re.match(r"(?i)^section\s+\d+[:.\-]\s+", ln):
        return False
    if re.match(rf"^\s*{QUESTION_ID_PATTERN}\.\s+", ln, flags=re.I):
        return True
    return False


def _parse_question_heading(line: str) -> tuple[str, str] | None:
    match = re.match(rf"^\s*({QUESTION_ID_PATTERN})\.\s+(.*)$", line or "", flags=re.I)
    if not match:
        return None
    prompt_id = match.group(1).strip()
    prompt_text = match.group(2).strip()
    if not prompt_text:
        return None
    return prompt_id, prompt_text


def _extract_prompt_options(details: str) -> list[str]:
    options: list[str] = []
    for raw_line in (details or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        normalized = line.lstrip("?").lstrip("•").lstrip("-").strip()
        if raw_line.strip().startswith(("?", "•", "-")) and normalized:
            options.append(normalized[:140])
    deduped: list[str] = []
    seen: set[str] = set()
    for option in options:
        key = option.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(option)
    return deduped[:10]


def _classify_prompt_type(prompt_text: str, details: str) -> str:
    combined = f"{prompt_text}\n{details}".lower()
    if "upload a file" in combined or "acceptable file types" in combined:
        return "upload"
    if re.search(r"\btopic\(s\)\b|\bcheck all that apply\b|\bselect (?:the )?topic", combined):
        return "multi_select"
    if re.search(
        r"\bwhich of the following\b|\bselect the geographic scope\b|\bselect the project stage\b|\bselect one\b|\bselect only one\b",
        combined,
    ):
        return "select_one"
    if re.search(r"\bselect one\b|\bselect only one\b", combined):
        return "select_one"
    if re.search(r"\bcheck all that apply\b|\bselect the option\(s\)\b", combined):
        return "multi_select"
    if re.search(r"\byes\b", combined) and re.search(r"\bno\b", combined):
        return "yes_no"
    if re.search(
        r"\bwebsite\b|\bemail\b|\bname\b|\bjob title\b|\btimezone\b|\bfounded\b|\bcurrent annual budget\b|\bfull-time employees\b|\benter the\b|\bplease enter\b",
        combined,
    ):
        return "short_field"
    if re.search(r"\bbudget category\b|\bbudget allocation\b|\bfunding request\b", combined):
        return "budget_row"
    if re.search(r"\bmilestone\b|\btimeframe\b|\boutcomes/key milestones\b", combined):
        return "milestone_row"
    return "narrative"


def _infer_response_style(prompt_text: str, details: str, prompt_type: str, section_title: str = "") -> str:
    combined = f"{section_title}\n{prompt_text}\n{details}".lower()
    word_limit = _extract_word_limit(f"{prompt_text}\n{details}")

    if re.search(r"\b(certification|certify|agree to the terms|privacy policy|authorized representative)\b", combined):
        return "fixed_attestation"
    if prompt_type == "upload":
        return "upload_placeholder"
    if prompt_type == "budget_row":
        return "budget_row"
    if prompt_type == "milestone_row":
        return "milestone_row"
    if prompt_type == "short_field":
        return "field"
    if prompt_type == "select_one":
        return "selection"
    if prompt_type == "multi_select":
        return "multi_selection"
    if prompt_type == "yes_no":
        return "yes_no_explanation"
    if isinstance(word_limit, int) and word_limit > 0 and word_limit <= 75:
        return "narrative_short"
    if isinstance(word_limit, int) and word_limit > 75:
        return "narrative_long"
    return "narrative_short"


def _infer_conditional_trigger(prompt_text: str) -> str | None:
    lowered = (prompt_text or "").lower()
    if not lowered.startswith("if "):
        return None
    if re.search(r'\bif (?:selected|you selected)\s+[\"“]?other[\"”]?', lowered):
        return "other"
    if re.search(r'\bif (?:answered|you answered|selected)\s+[\"“]?yes[\"”]?', lowered) or re.search(r"^if yes\b|^if you answered yes\b|^if the answer is yes\b", lowered):
        return "yes"
    if re.search(r'\bif (?:answered|you answered|selected)\s+[\"“]?no[\"”]?', lowered) or re.search(r"^if no\b|^if the answer is no\b", lowered):
        return "no"
    return "conditional"


def _extract_sections_from_question_template(lines: list[str]) -> list[dict]:
    sections: list[dict] = []
    current_section: dict[str, Any] | None = None
    current_prompt: dict[str, Any] | None = None

    def flush_prompt() -> None:
        nonlocal current_prompt, current_section
        if not current_prompt or not current_section:
            current_prompt = None
            return
        prompt_text = str(current_prompt.get("prompt_text") or "").strip()
        details = "\n".join(current_prompt.get("detail_lines", [])).strip()
        if details:
            current_prompt["detail_text"] = details[:2000]
        prompt_type = _classify_prompt_type(prompt_text, details)
        response_style = _infer_response_style(
            prompt_text,
            details,
            prompt_type,
            str(current_section.get("title") or ""),
        )
        options = _extract_prompt_options(details)
        conditional_trigger = _infer_conditional_trigger(prompt_text)
        current_prompt["word_limit"] = _extract_word_limit(f"{prompt_text}\n{details}")
        prompt_item = {
            "prompt_id": current_prompt.get("prompt_id"),
            "label": current_prompt.get("prompt_id"),
            "prompt_text": prompt_text,
            "prompt_type": prompt_type,
            "response_style": response_style,
            "answer_type": _answer_type_from_response_style(response_style),
            "word_limit": current_prompt.get("word_limit"),
            "detail_text": current_prompt.get("detail_text", ""),
            "required": _is_prompt_required(prompt_text, details),
            "sub_prompt": _is_sub_prompt_id(str(current_prompt.get("prompt_id") or "")),
            "options": options,
        }
        if conditional_trigger:
            existing_items = current_section.get("prompt_items", [])
            if existing_items:
                prompt_item["parent_prompt_id"] = existing_items[-1].get("prompt_id")
                prompt_item["conditional_on_previous"] = conditional_trigger
        prompt_item = _normalize_prompt_item_schema(
            prompt_item,
            str(current_section.get("key") or ""),
        )
        current_section.setdefault("prompt_items", []).append(prompt_item)
        current_prompt = None

    def flush_section() -> None:
        nonlocal current_section
        flush_prompt()
        if not current_section:
            return
        prompt_items = current_section.get("prompt_items", [])
        if prompt_items:
            current_section["guidance"] = "\n".join(
                str(item.get("prompt_text") or "").strip()
                for item in prompt_items
                if str(item.get("prompt_text") or "").strip()
            )[:3000]
        sections.append(current_section)
        current_section = None

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if re.match(r"(?i)^section\s+\d+[:.\-]\s+", line):
            flush_section()
            current_section = {
                "key": _normalize_section_key(line, len(sections)),
                "title": line,
                "guidance": "",
                "prompt_items": [],
            }
            continue
        if _is_question_heading(line):
            if current_section is None:
                current_section = {
                    "key": _normalize_section_key(f"Section {len(sections)+1}", len(sections)),
                    "title": f"Section {len(sections)+1}",
                    "guidance": "",
                    "prompt_items": [],
                }
            flush_prompt()
            parsed = _parse_question_heading(line)
            prompt_id = (parsed[0] if parsed else f"Q{len(current_section.get('prompt_items', []))+1}").strip()
            prompt_text = (parsed[1] if parsed else line).strip()
            current_prompt = {
                "prompt_id": prompt_id,
                "prompt_text": prompt_text.rstrip(),
                "detail_lines": [],
            }
            continue
        if current_prompt is not None:
            if line.upper() in {"RESPONSE", "YES", "NO"}:
                continue
            if re.match(r"(?i)^google\.org impact challenge|^application template\b|^modified:\s+.+page\s+\d+\s+of\s+\d+", line):
                continue
            if re.match(r"(?i)^(first|last) name\s*:|^email\b|^website\b|^utc\s*-?\d+", line):
                current_prompt["detail_lines"].append(line)
                continue
            if re.match(r"(?i)^limit\b|^limit:\s*\d+", line) or _extract_word_limit(line):
                current_prompt["detail_lines"].append(line)
                continue
            if line.startswith(("-", "*")) or re.match(r"^[A-Z][a-z]+\b", line):
                current_prompt["detail_lines"].append(line)
                continue
            current_prompt["detail_lines"].append(line)

    flush_section()
    return _normalize_sections(sections)


def _extract_word_limit(text: str) -> int | None:
    m = re.search(
        r"(?i)\b(?:word\s*limit|max(?:imum)?\s*words?|up to)\s*[:\-]?\s*(\d{2,5})\b",
        text,
    )
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None

    m = re.search(r"(?i)\b(\d{2,5})\s*words?\b", text)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None
    return None


def _is_prompt_required(prompt_text: str, details: str) -> bool:
    combined = f"{prompt_text}\n{details}"
    return "*" in combined or bool(re.search(r"(?i)\brequired\b", combined))


def _is_sub_prompt_id(prompt_id: str) -> bool:
    token = (prompt_id or "").strip().lower()
    return bool(re.search(r"[a-z]$|\.", token))


def _answer_type_from_response_style(response_style: str) -> str:
    mapping = {
        "field": "field",
        "selection": "selection",
        "multi_selection": "multi_selection",
        "yes_no_explanation": "yes_no",
        "narrative_short": "narrative_short",
        "narrative_long": "narrative_long",
        "budget_row": "budget_row",
        "milestone_row": "milestone_row",
        "upload_placeholder": "upload",
        "fixed_attestation": "attestation",
    }
    return mapping.get(response_style or "", "narrative_short")


def _normalize_prompt_item_schema(prompt_item: dict[str, Any], section_key: str) -> dict[str, Any]:
    prompt_id = str(prompt_item.get("prompt_id") or "").strip()
    prompt_text = str(prompt_item.get("prompt_text") or "").strip()
    detail_text = str(prompt_item.get("detail_text") or "").strip()
    prompt_type = str(prompt_item.get("prompt_type") or "narrative")
    response_style = str(prompt_item.get("response_style") or "")
    normalized = {
        "prompt_id": prompt_id,
        "label": prompt_id,
        "section_id": section_key,
        "prompt_text": prompt_text,
        "detail_text": detail_text,
        "prompt_type": prompt_type,
        "response_style": response_style,
        "answer_type": _answer_type_from_response_style(response_style),
        "word_limit": prompt_item.get("word_limit"),
        "required": bool(prompt_item.get("required")) or _is_prompt_required(prompt_text, detail_text),
        "sub_prompt": _is_sub_prompt_id(prompt_id),
        "options": prompt_item.get("options") or [],
        "conditional_on_previous": prompt_item.get("conditional_on_previous"),
        "parent_prompt_id": prompt_item.get("parent_prompt_id"),
        "response_value": prompt_item.get("response_value"),
        "source_confidence": prompt_item.get("source_confidence"),
        "source_origin": prompt_item.get("source_origin"),
    }
    return normalized


def _section_question_signal_count(text: str) -> int:
    if not text:
        return 0
    return len(
        re.findall(
            rf"(?im)^\s*(?:{QUESTION_ID_PATTERN})\.\s+|^\s*(?:describe|explain|provide|list|identify|select|please enter)\b",
            text,
        )
    )


def _prompt_map_quality_issues(section: dict[str, Any]) -> list[str]:
    issues: list[str] = []
    prompt_items = section.get("prompt_items") or []
    guidance = str(section.get("guidance") or "")
    signal_count = _section_question_signal_count(guidance)
    if signal_count >= 3 and len(prompt_items) <= 1:
        issues.append("few_prompts_extracted_relative_to_question_signals")

    prompt_texts = [str(item.get("prompt_text") or "").strip().lower() for item in prompt_items]
    if prompt_texts and len(set(prompt_texts)) < len(prompt_texts):
        issues.append("duplicate_prompt_texts_detected")

    if sum(1 for text in prompt_texts if len(text) > 180) >= 2:
        issues.append("multiple_prompt_texts_are_overlong")

    if sum(1 for item in prompt_items if not item.get("answer_type")):
        issues.append("missing_answer_types")
    return issues


def _should_normalize_prompt_items_with_llm(section: dict[str, Any]) -> bool:
    issues = _prompt_map_quality_issues(section)
    if not issues:
        return False
    guidance = str(section.get("guidance") or "")
    if len(guidance) < 300:
        return False
    return True


def _normalize_prompt_items_with_llm(section: dict[str, Any]) -> tuple[list[dict[str, Any]], str | None]:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return [], "missing_openai_api_key"

    payload = {
        "task": "Normalize extracted application prompts for one section.",
        "section": {
            "title": section.get("title"),
            "guidance": section.get("guidance"),
            "existing_prompt_items": section.get("prompt_items", []),
        },
        "rules": [
            "Return JSON only.",
            "Extract the actual applicant response prompts and sub-prompts from this section.",
            "Preserve prompt order.",
            "Keep wording close to the source text; do not paraphrase unless needed to remove obvious OCR noise.",
            "Infer answer_type from the prompt semantics using only: field, selection, multi_selection, yes_no, narrative_short, narrative_long, budget_row, milestone_row, upload, attestation.",
            "Infer required=true only when explicit.",
            "If a word limit is explicit, include it; otherwise omit it.",
            "Use parent_prompt_id only for clear conditional follow-ups.",
        ],
        "schema": {
            "prompt_items": [
                {
                    "prompt_id": "q16a",
                    "label": "Q16a",
                    "prompt_text": "The public service challenge this project addresses is ...",
                    "detail_text": "Limit: 100 words",
                    "word_limit": 100,
                    "required": True,
                    "sub_prompt": True,
                    "answer_type": "narrative_short",
                    "options": [],
                    "parent_prompt_id": None,
                    "conditional_on_previous": None,
                }
            ]
        },
    }

    try:
        from openai import OpenAI

        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "You normalize application prompts into strict JSON. You do extraction only, not drafting.",
                },
                {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
            ],
            response_format={"type": "json_object"},
            temperature=0.0,
        )
        data = json.loads(resp.choices[0].message.content or "{}")
    except Exception as exc:
        logger.exception("Prompt normalization LLM pass failed")
        return [], f"{type(exc).__name__}: {exc}"

    prompt_items: list[dict[str, Any]] = []
    for item in data.get("prompt_items") or []:
        prompt_id = str(item.get("prompt_id") or item.get("label") or "").strip()
        prompt_text = str(item.get("prompt_text") or "").strip()
        if not prompt_id or not prompt_text:
            continue
        normalized = {
            "prompt_id": prompt_id,
            "label": str(item.get("label") or prompt_id).strip(),
            "section_id": str(section.get("key") or ""),
            "prompt_text": prompt_text,
            "detail_text": str(item.get("detail_text") or "").strip()[:2000],
            "prompt_type": str(item.get("prompt_type") or "narrative"),
            "response_style": str(item.get("response_style") or ""),
            "answer_type": str(item.get("answer_type") or "narrative_short"),
            "word_limit": item.get("word_limit"),
            "required": bool(item.get("required")),
            "sub_prompt": bool(item.get("sub_prompt")) or _is_sub_prompt_id(prompt_id),
            "options": item.get("options") or [],
            "conditional_on_previous": item.get("conditional_on_previous"),
            "parent_prompt_id": item.get("parent_prompt_id"),
        }
        prompt_items.append(normalized)
    return prompt_items, None

def _normalize_section_key(title: str, index: int) -> str:
    key_base = re.sub(r"[^a-z0-9]+", "_", (title or "").lower()).strip("_")[:60]
    return key_base or f"section_{index + 1}"


def _infer_section_purpose(title: str, guidance: str, prompt_items: list[dict]) -> str:
    haystack = f"{title}\n{guidance}".lower()
    prompt_types = [str(item.get("prompt_type") or "narrative") for item in (prompt_items or [])]
    prompt_count = len(prompt_types)

    if re.search(r"\b(certification|attestation|compliance|ethics|privacy|consent|declaration|disclosure|conflict of interest)\b", haystack):
        return "compliance_attestation"

    if re.search(r"\b(budget|timeline|milestone|workplan|schedule|allocation|funding request)\b", haystack):
        return "budget_timeline"

    if prompt_count:
        budget_like = sum(1 for prompt_type in prompt_types if prompt_type in {"budget_row", "milestone_row"})
        if budget_like >= max(1, prompt_count // 3):
            return "budget_timeline"

        attachment_like = sum(1 for prompt_type in prompt_types if prompt_type == "upload")
        if attachment_like >= max(1, prompt_count // 2):
            return "attachments"

        structured_like = sum(
            1 for prompt_type in prompt_types
            if prompt_type in {"short_field", "select_one", "multi_select", "yes_no"}
        )
        if structured_like >= max(1, (prompt_count + 1) // 2):
            return "structured_form"

    if re.search(r"\b(organization|submitter|applicant|contact|profile|eligibility)\b", haystack):
        return "structured_form"

    if re.search(r"\b(attachment|supporting document|upload|appendix|appendices|supplemental)\b", haystack):
        return "attachments"

    return "narrative"


def _prompt_semantic_bucket(prompt_item: dict[str, Any]) -> str:
    response_style = str(prompt_item.get("response_style") or "")
    prompt_type = str(prompt_item.get("prompt_type") or "")
    prompt_text = str(prompt_item.get("prompt_text") or "").lower()
    detail_text = str(prompt_item.get("detail_text") or "").lower()
    haystack = f"{prompt_text}\n{detail_text}"

    if response_style in {"fixed_attestation", "yes_no_explanation"} or re.search(
        r"\b(certification|certify|attestation|compliance|ethics|privacy|consent|declaration|disclosure|conflict of interest)\b",
        haystack,
    ):
        return "compliance_attestation"
    if response_style in {"budget_row", "milestone_row"} or prompt_type in {"budget_row", "milestone_row"} or re.search(
        r"\b(budget|allocation|funding request|milestone|timeline|timeframe|workplan|schedule)\b",
        haystack,
    ):
        return "budget_timeline"
    if response_style in {"field", "selection", "multi_selection", "upload_placeholder"} or prompt_type in {
        "short_field",
        "select_one",
        "multi_select",
        "upload",
    }:
        return "structured_form"
    return "narrative"


def _repair_section_prompt_assignment(sections: list[dict]) -> list[dict]:
    if not sections:
        return sections

    target_sections: dict[str, list[dict]] = {}
    for section in sections:
        purpose = str(section.get("section_purpose") or "narrative")
        target_sections.setdefault(purpose, []).append(section)

    repaired_sections: list[dict] = []
    displaced_prompts: dict[str, list[dict]] = {str(section.get("key") or ""): [] for section in sections}

    for index, section in enumerate(sections):
        current_key = str(section.get("key") or "")
        current_purpose = str(section.get("section_purpose") or "narrative")
        prompt_items = section.get("prompt_items") or []
        kept_prompts: list[dict] = []

        for prompt_item in prompt_items:
            bucket = _prompt_semantic_bucket(prompt_item)
            if bucket == current_purpose or current_purpose in {"narrative", "attachments", "structured_form"}:
                kept_prompts.append(prompt_item)
                continue

            # Keep this repair narrow: only pull obviously misplaced prompts out of
            # compliance/certification and budget/timeline sections.
            if current_purpose not in {"compliance_attestation", "budget_timeline"}:
                kept_prompts.append(prompt_item)
                continue

            if current_purpose == "compliance_attestation" and bucket != "budget_timeline":
                kept_prompts.append(prompt_item)
                continue
            if current_purpose == "budget_timeline" and bucket not in {"compliance_attestation"}:
                kept_prompts.append(prompt_item)
                continue

            candidates = target_sections.get(bucket) or []
            if not candidates:
                kept_prompts.append(prompt_item)
                continue

            nearest = min(
                candidates,
                key=lambda candidate: abs(
                    sections.index(candidate) - index
                ),
            )
            displaced_prompts[str(nearest.get("key") or "")].append(prompt_item)

        section = dict(section)
        section["prompt_items"] = kept_prompts
        if len(kept_prompts) != len(prompt_items):
            section.setdefault("parser_diagnostics", []).append("prompt_reassignment_applied")
        repaired_sections.append(section)

    final_sections: list[dict] = []
    for section in repaired_sections:
        key = str(section.get("key") or "")
        appended = displaced_prompts.get(key) or []
        if appended:
            merged = list(section.get("prompt_items") or []) + appended
            section["prompt_items"] = merged
            guidance_parts = [str(item.get("prompt_text") or "").strip() for item in merged if str(item.get("prompt_text") or "").strip()]
            section["guidance"] = "\n".join(guidance_parts)[:3000]
            section.setdefault("parser_diagnostics", []).append("prompt_reassignment_received")
        final_sections.append(section)

    return final_sections


def _extract_question_prompts(guidance: str, max_prompts: int = 25) -> list[dict]:
    prompts: list[dict] = []
    seen: set[str] = set()
    trigger = re.compile(
        r"^(?:[-*•]\s*|\(?[a-z0-9]{1,3}[\).]\s+)?"
        r"(describe|explain|provide|outline|summarize|identify|demonstrate|detail|list|state|tell us|please describe)\b",
        flags=re.I,
    )
    for raw_line in (guidance or "").splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip(" -•\t")
        if not line or len(line) < 12:
            continue
        is_prompt = line.endswith("?") or bool(trigger.match(line))
        if not is_prompt:
            continue
        for prompt_text in _expand_compound_prompt(line):
            normalized = prompt_text.lower().rstrip(".?")
            if normalized in seen:
                continue
            seen.add(normalized)
            prompt_type = "question" if prompt_text.endswith("?") else "instruction"
            prompt_word_limit = _extract_word_limit(prompt_text)
            detail_text = ""
            classified_type = _classify_prompt_type(prompt_text, detail_text)
            prompts.append(
                _normalize_prompt_item_schema(
                    {
                        "prompt_id": f"prompt_{len(prompts) + 1}",
                        "label": f"prompt_{len(prompts) + 1}",
                        "prompt_text": prompt_text.rstrip(),
                        "prompt_type": classified_type if classified_type != "narrative" else prompt_type,
                        "response_style": _infer_response_style(
                            prompt_text,
                            detail_text,
                            classified_type if classified_type != "narrative" else prompt_type,
                        ),
                        "word_limit": prompt_word_limit,
                        "detail_text": detail_text,
                        "options": [],
                    },
                    "",
                )
            )
            if len(prompts) >= max_prompts:
                break
        if len(prompts) >= max_prompts:
            break
    return prompts


def _summarize_section_guidance(guidance: str, prompt_items: list[dict]) -> str:
    prompt_texts = [
        str(item.get("prompt_text") or "").strip()
        for item in (prompt_items or [])
        if str(item.get("prompt_text") or "").strip()
    ]
    if prompt_texts:
        return "\n".join(prompt_texts[:12])[:3000]
    return str(guidance or "").strip()[:3000]


def _expand_compound_prompt(line: str) -> list[str]:
    cleaned = re.sub(r"\s+", " ", (line or "").strip()).rstrip(".")
    if not cleaned:
        return []

    expanded: list[str] = []

    verb_splitter = re.compile(
        r"\s+(?:and|as well as)\s+(?=(?:describe|explain|provide|outline|summarize|identify|demonstrate|detail|list|state)\b)",
        flags=re.I,
    )
    verb_parts = [part.strip(" ,") for part in verb_splitter.split(cleaned) if part.strip(" ,")]
    if len(verb_parts) > 1:
        for part in verb_parts:
            expanded.extend(_expand_compound_prompt(part))
        return expanded

    including_match = re.search(r"(?i)\b(including|covering|addressing)\b\s+(.*)$", cleaned)
    if including_match:
        prefix = cleaned[: including_match.start()].rstrip(" ,")
        suffix = including_match.group(2).strip()
        if prefix:
            expanded.append(prefix if prefix.endswith("?") else f"{prefix}.")
        list_items = _split_prompt_list(suffix)
        subject = _prompt_subject(prefix)
        lead_verb = _prompt_verb(prefix) or "Describe"
        for item in list_items:
            item = item.strip(" .")
            if not item:
                continue
            prompt = f"{lead_verb} {subject} {item}.".replace("  ", " ").strip()
            expanded.append(prompt)
        return _dedupe_prompt_list(expanded)

    comma_match = re.match(r"(?i)^(describe|explain|provide|outline|summarize|identify|demonstrate|detail|list|state)\s+([^,]+),\s*(.+)$", cleaned)
    if comma_match:
        verb = comma_match.group(1).capitalize()
        subject = comma_match.group(2).strip()
        remainder = comma_match.group(3).strip()
        subject, seed_item = _split_subject_seed(subject)
        if seed_item:
            expanded.append(f"{verb} {subject} {seed_item}.".replace("  ", " ").strip())
        else:
            expanded.append(f"{verb} {subject}.")
        for item in _split_prompt_list(remainder):
            item = item.strip(" .")
            if not item:
                continue
            expanded.append(_compose_prompt_with_subject(verb, subject, item))
        return _dedupe_prompt_list(expanded)

    return [cleaned if cleaned.endswith("?") else f"{cleaned}."]


def _split_prompt_list(text: str) -> list[str]:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    normalized = re.sub(r"\s*,\s*and\s+", ", ", normalized, flags=re.I)
    normalized = re.sub(r"\s+and\s+", ", ", normalized, flags=re.I)
    parts = [part.strip(" ,.") for part in normalized.split(",") if part.strip(" ,.")]
    return parts[:6]


def _split_subject_seed(subject: str) -> tuple[str, str | None]:
    cleaned = re.sub(r"\s+", " ", subject or "").strip(" ,.")
    if not cleaned:
        return "the project", None
    match = re.match(r"^(.*?\b(?:its|their|the project'?s|your organization'?s|the organization'?s))\s+([a-z][a-z0-9 /-]*)$", cleaned, flags=re.I)
    if match:
        return match.group(1).strip(), match.group(2).strip()
    if cleaned.endswith("'s"):
        return cleaned, None
    return cleaned, None


def _compose_prompt_with_subject(verb: str, subject: str, item: str) -> str:
    item = item.strip(" .")
    if not item:
        return f"{verb} {subject}."
    lowered_subject = subject.lower()
    if re.match(r"^(its|their|the project'?s|your organization'?s|the organization'?s)\b", item, flags=re.I):
        prompt = f"{verb} {subject} {item}."
    elif lowered_subject.endswith("'s") or lowered_subject.endswith("its") or lowered_subject.endswith("their"):
        prompt = f"{verb} {subject} {item}."
    else:
        prompt = f"{verb} {subject} {item}."
    return re.sub(r"\s+", " ", prompt).strip()


def _prompt_subject(prefix: str) -> str:
    cleaned = re.sub(
        r"(?i)^(describe|explain|provide|outline|summarize|identify|demonstrate|detail|list|state)\s+",
        "",
        prefix or "",
    ).strip(" ,")
    return cleaned or "the project"


def _prompt_verb(prefix: str) -> str | None:
    match = re.match(
        r"(?i)^(describe|explain|provide|outline|summarize|identify|demonstrate|detail|list|state)\b",
        prefix or "",
    )
    if not match:
        return None
    return match.group(1).capitalize()


def _dedupe_prompt_list(prompts: list[str]) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for prompt in prompts:
        normalized = re.sub(r"\s+", " ", prompt.lower()).strip(" .?")
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        out.append(prompt)
    return out

def _normalize_sections(sections: list[dict]) -> list[dict]:
    out: list[dict] = []
    seen: dict[str, int] = {}

    for idx, sec in enumerate(sections):
        title = str(sec.get("title") or "").strip()
        guidance = str(sec.get("guidance") or "").strip()
        if not title:
            continue

        base_key = str(sec.get("key") or _normalize_section_key(title, idx)).strip()
        key = re.sub(r"[^a-z0-9_]+", "_", base_key.lower()).strip("_") or _normalize_section_key(title, idx)
        cnt = seen.get(key, 0) + 1
        seen[key] = cnt
        key = key if cnt == 1 else f"{key}_{cnt}"

        item: dict[str, Any] = {
            "key": key,
            "title": title[:140],
            "guidance": guidance[:3000],
        }
        prompt_items = sec.get("prompt_items")
        if isinstance(prompt_items, list) and prompt_items:
            item["prompt_items"] = prompt_items
        else:
            inferred_prompts = _extract_question_prompts(guidance)
            if inferred_prompts:
                item["prompt_items"] = inferred_prompts

        prompt_items_out = [
            _normalize_prompt_item_schema(prompt_item, key)
            for prompt_item in item.get("prompt_items", [])
        ]
        item["prompt_items"] = prompt_items_out
        item["guidance"] = _summarize_section_guidance(guidance, prompt_items_out)

        section_word_limit = _resolve_section_word_limit(
            title=title,
            guidance=guidance,
            prompt_items=prompt_items_out,
            explicit_word_limit=sec.get("word_limit"),
        )
        if section_word_limit:
            item["word_limit"] = section_word_limit

        item["section_purpose"] = str(
            sec.get("section_purpose")
            or _infer_section_purpose(title, guidance, prompt_items_out)
        )

        out.append(item)

    return out


def _resolve_section_word_limit(
    *,
    title: str,
    guidance: str,
    prompt_items: list[dict],
    explicit_word_limit: Any,
) -> int | None:
    if isinstance(explicit_word_limit, int) and explicit_word_limit > 0:
        candidate = explicit_word_limit
    else:
        candidate = _extract_word_limit(f"{title}\n{guidance}")

    if not candidate:
        return None

    if not prompt_items or len(prompt_items) <= 1:
        return candidate

    scope_text = f"{title}\n{guidance}".lower()
    if re.search(r"\b(this section|entire section|overall response|total response|full response|combined response|in total)\b", scope_text):
        return candidate

    # Multiple prompt items with no explicit section-scoped wording: avoid showing
    # a misleading blanket limit at the section level. Prompt-level limits remain available.
    return None

def _should_use_llm_fallback(raw_text: str, heuristic_sections: list[dict]) -> tuple[bool, list[str]]:
    reasons: list[str] = []
    if len(raw_text.strip()) < 800:
        return False, reasons
    if len(heuristic_sections) <= 1:
        reasons.append("heuristic_detected_1_or_fewer_sections")
    if len(heuristic_sections) > 12:
        reasons.append("heuristic_detected_more_than_12_sections")

    titles = [str(s.get("title") or "").strip() for s in heuristic_sections]
    if titles:
        very_short = sum(1 for t in titles if len(t.split()) <= 2)
        if len(titles) >= 4 and (very_short / len(titles)) >= 0.5:
            reasons.append("many_heuristic_titles_are_very_short")

        normalized_titles = [
            re.sub(r"\d+", "", re.sub(r"[^a-z0-9]+", "", t.lower()))
            for t in titles
            if t
        ]
        if normalized_titles:
            unique_count = len(set(normalized_titles))
            if unique_count <= max(2, len(normalized_titles) // 3):
                reasons.append("heuristic_titles_have_low_uniqueness")

    keys = {str(s.get("key") or "") for s in heuristic_sections}
    if keys == {"application_requirements"}:
        reasons.append("heuristic_collapsed_to_application_requirements")

    return bool(reasons), reasons


def _extract_sections_with_llm(text: str) -> tuple[list[dict], str | None]:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return [], "missing_openai_api_key"

    snippet = text[:20000]
    if len(snippet.strip()) < 800:
        return [], "insufficient_text_for_llm_fallback"

    prompt = {
        "task": "Extract concrete grant application sections from the posting text.",
        "rules": [
            "Return JSON only.",
            "Prefer explicit required/expected sections in the application package.",
            "Prefer top-level numbered headings (Roman numerals, numeric levels, lettered lists) when present.",
            "Return sections the applicant must respond to; exclude context-only headings and administrative notes.",
            "Do not split a single required heading into multiple micro-sections unless the document clearly requires separate responses.",
            "Do not invent sections not implied by the text.",
            "Each section needs key, title, guidance, and optional word_limit.",
            "If limits are not explicit, omit word_limit.",
        ],
        "schema": {
            "sections": [
                {
                    "key": "snake_case_key",
                    "title": "Section Title",
                    "guidance": "What the applicant should cover in this section",
                    "word_limit": 500,
                }
            ]
        },
        "text": snippet,
    }

    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=CHAT_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You extract grant requirement sections from messy PDF text. "
                        "Output strict JSON with high recall and minimal hallucination."
                    ),
                },
                {"role": "user", "content": json.dumps(prompt, ensure_ascii=False)},
            ],
            response_format={"type": "json_object"},
            temperature=0.0,
        )
        data = json.loads(resp.choices[0].message.content or "{}")
        return _normalize_sections((data.get("sections") or [])), None
    except Exception as exc:
        logger.exception("Grant section LLM fallback failed")
        return [], f"{type(exc).__name__}: {exc}"

def _extract_sections_from_text(text: str) -> list[dict]:
    """
    Heuristic section extraction for grant packages.
    Finds heading-like lines and assigns following content as guidance.
    If no clear headings are found, returns a single catch-all section.
    """
    cleaned = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not cleaned:
        return [{"key": "application_requirements", "title": "Application Requirements", "guidance": ""}]

    lines = [ln.strip() for ln in cleaned.splitlines()]
    question_template_sections = _extract_sections_from_question_template(lines)
    if len(question_template_sections) >= 4:
        return question_template_sections
    numbered_sections = _extract_sections_from_numbered_headings(lines)
    if numbered_sections:
        return numbered_sections

    heading_idxs = []
    for i, ln in enumerate(lines):
        if _is_probable_heading(ln):
            # Avoid duplicate heading candidates on adjacent lines.
            if heading_idxs and i - heading_idxs[-1] <= 1 and lines[heading_idxs[-1]].lower() == ln.lower():
                continue
            heading_idxs.append(i)

    # If we cannot confidently segment headings, keep all guidance in one section
    # instead of forcing a fixed template.
    if len(heading_idxs) < 2:
        return [
            {
                "key": "application_requirements",
                "title": "Application Requirements",
                "guidance": cleaned[:6000],
                "word_limit": _extract_word_limit(cleaned),
            }
        ]

    sections = []
    seen_keys: dict[str, int] = {}
    for idx, start in enumerate(heading_idxs):
        end = heading_idxs[idx + 1] if idx + 1 < len(heading_idxs) else len(lines)
        heading = lines[start].rstrip(":").strip()
        body = "\n".join(lines[start + 1:end]).strip()
        key_base = re.sub(r"[^a-z0-9]+", "_", heading.lower()).strip("_")[:60] or f"section_{idx+1}"
        count = seen_keys.get(key_base, 0) + 1
        seen_keys[key_base] = count
        key = key_base if count == 1 else f"{key_base}_{count}"

        # Keep enough context to guide generation, but cap to avoid runaway payloads.
        guidance = body[:3000]
        sec = {"key": key, "title": heading, "guidance": guidance}
        wl = _extract_word_limit(f"{heading}\n{body}")
        if wl:
            sec["word_limit"] = wl
        sections.append(sec)

    return _normalize_sections(sections)

def parse_grant_upload_to_requirements(uploaded_file) -> Tuple[Dict[str, Any] | None, str]:
    """
    Returns:
      requirements: dict usable by the app
      raw_text: extracted text from the document (best-effort)
    """
    name = uploaded_file.name.lower()

    file_bytes = _read_uploaded_file_bytes(uploaded_file)

    if name.endswith(".txt"):
        raw = _read_txt(uploaded_file)
    elif name.endswith(".pdf"):
        raw = _read_pdf(uploaded_file)
    elif name.endswith(".docx"):
        raw = _read_docx(uploaded_file)
    else:
        raw = ""

    raw = (raw or "").strip()
    heuristic_sections = _extract_sections_from_text(raw)
    sections = heuristic_sections
    parser_mode = "heuristic"
    llm_used = False
    docai_used = False
    docai_error: str | None = None
    docai_meta: dict[str, Any] = {}
    llm_error: str | None = None
    fallback_reasons: list[str] = []
    diagnostics: list[str] = []

    if name.endswith(".pdf") and _document_ai_is_configured():
        docai_sections, docai_raw, docai_meta, docai_error = _extract_sections_with_document_ai(
            file_bytes,
            "application/pdf",
        )
        docai_used = True
        if docai_raw and len(docai_raw) >= len(raw):
            raw = docai_raw.strip()
        if len(docai_sections) >= 2:
            heuristic_sections = docai_sections
            sections = docai_sections
            parser_mode = "document_ai"
            diagnostics.append("document_ai_selected")
        else:
            diagnostics.append("document_ai_did_not_return_enough_sections")

    should_use_llm, fallback_reasons = _should_use_llm_fallback(raw, heuristic_sections)
    diagnostics.extend(fallback_reasons)

    if should_use_llm:
        llm_sections, llm_error = _extract_sections_with_llm(raw)
        llm_used = True
        if len(llm_sections) >= 2:
            sections = llm_sections
            parser_mode = "llm_fallback"
            diagnostics.append("llm_fallback_selected")
        else:
            parser_mode = "heuristic_fallback_single"
            diagnostics.append("llm_fallback_did_not_return_enough_sections")

    confidence = "high" if len(sections) >= 4 else ("medium" if len(sections) >= 2 else "low")
    section_titles_preview = [str(s.get("title") or "") for s in sections[:8]]
    heuristic_titles_preview = [str(s.get("title") or "") for s in heuristic_sections[:8]]

    logger.info(
        "Grant parser completed file=%s mode=%s confidence=%s raw_text_length=%s heuristic_sections=%s final_sections=%s docai_used=%s llm_used=%s reasons=%s docai_error=%s llm_error=%s",
        uploaded_file.name,
        parser_mode,
        confidence,
        len(raw),
        len(heuristic_sections),
        len(sections),
        docai_used,
        llm_used,
        ",".join(fallback_reasons) or "none",
        docai_error or "none",
        llm_error or "none",
    )

    # Minimal requirements schema the rest of the app expects
    requirements = {
        "grant_name": uploaded_file.name,
        "sections": sections,
        "eligibility": [],            # optional / can be expanded later
        "word_limits": {},            # optional
        "must_include": [],           # optional
        "raw_text": raw,
        "required_sections": [s.get("title") for s in sections],
        "parser_meta": {
            "mode": parser_mode,
            "confidence": confidence,
            "model": CHAT_MODEL,
            "raw_text_length": len(raw),
            "heuristic_section_count": len(heuristic_sections),
            "final_section_count": len(sections),
            "document_ai_used": docai_used,
            "document_ai_error": docai_error,
            "document_ai_page_count": docai_meta.get("page_count"),
            "document_ai_form_field_count": docai_meta.get("form_field_count"),
            "document_ai_table_count": docai_meta.get("table_count"),
            "document_ai_used_augmented_text": docai_meta.get("used_augmented_text"),
            "document_ai_location": docai_meta.get("location"),
            "llm_fallback_used": llm_used,
            "llm_error": llm_error,
            "fallback_reasons": fallback_reasons,
            "diagnostics": diagnostics,
            "heuristic_titles_preview": heuristic_titles_preview,
            "section_titles_preview": section_titles_preview,
            "question_count": sum(len(s.get("prompt_items", [])) for s in sections),
            "structured_prompt_count": sum(
                sum(1 for prompt in (s.get("prompt_items") or []) if prompt.get("answer_type"))
                for s in sections
            ),
            "sections_with_prompt_quality_warnings": sum(
                1
                for s in sections
                if _prompt_map_quality_issues(s) or s.get("parser_diagnostics")
            ),
        },
    }

    return requirements, raw

