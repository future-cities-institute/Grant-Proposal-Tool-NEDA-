"""
FastAPI backend for the Grant Proposal Builder frontend.
Run from repo root: uvicorn api.main:app --reload --port 8000
Set PYTHONPATH to include repo root so backend.app imports resolve.
"""
from __future__ import annotations

import sys
import re
import os
import logging
from io import BytesIO
from pathlib import Path
from datetime import date

# Ensure repo root is on path (ced-suite)
ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# Load api/.env if present (local dev)
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env")
except ImportError:
    pass

from fastapi import Depends, FastAPI, File, Header, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from typing import Any, Dict, List, Optional

from backend.app.compliance import build_default_service, ComplianceEvaluationService
from backend.app.compliance.proposal_analysis import ProposalAnalysisService
from backend.app.compliance.proposal_rubric import RUBRIC_VERSION
from backend.app.compliance.models import (
    ComplianceEvaluationRequest,
    ComplianceEvaluationResponse,
)
from backend.app.compliance.proposal_models import (
    ProposalAnalysisResponse,
    ProposalChatRequest,
    ProposalChatResponse,
    ProposalReanalyzeRequest,
    ProposalSectionRewriteRequest,
    ProposalSectionRewriteResponse,
    ProposalSection,
)
from backend.app.auth import user_from_authorization_header
from backend.app.workspace_store import (
    create_feedback_report as store_create_feedback_report,
    create_proposal as store_create_proposal,
    delete_feedback_report as store_delete_feedback_report,
    delete_proposal as store_delete_proposal,
    duplicate_proposal as store_duplicate_proposal,
    get_or_create_user,
    get_community_profile as store_get_community_profile,
    get_feedback_report as store_get_feedback_report,
    get_proposal as store_get_proposal,
    init_workspace_store,
    list_proposals as store_list_proposals,
    list_feedback_reports as store_list_feedback_reports,
    mark_proposal_exported,
    update_proposal as store_update_proposal,
    upsert_community_profile as store_upsert_community_profile,
)

logger = logging.getLogger("uvicorn.error")

app = FastAPI(title="Grant Proposal API", version="0.1.0")

cors_origins = [
    origin.strip().rstrip("/")
    for origin in os.getenv("CORS_ORIGINS", "").split(",")
    if origin.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_origin_regex=r"^https?://((localhost|127\.0\.0\.1)(:\d+)?|[a-z0-9-]+\.vercel\.app)$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

init_workspace_store()


# ---------- Pydantic models for request/response ----------
class CommunityProfile(BaseModel):
    community_name: str = ""
    region: str = ""
    local_priority: str = ""
    legal_name: str = ""
    operating_name: str = ""
    applicant_profile: str = ""
    registration_number: str = ""
    year_established: str = ""
    contact_name: str = ""
    contact_title: str = ""
    contact_email: str = ""
    contact_phone: str = ""
    mailing_address: str = ""
    website: str = ""
    indigenous_communities: str = ""
    population_served: str = ""
    demographic_context: str = ""
    existing_services: str = ""
    service_gaps: str = ""
    remoteness_context: str = ""
    governance_context: str = ""
    project_title: str = ""
    project_location: str = ""
    timeline: str = ""
    challenges: str = ""
    strengths: str = ""
    partners: str = ""
    applicant_type: str = ""
    project_type: str = ""
    project_stage: str = ""
    community_support_status: str = ""
    other_funding_status: str = ""
    project_summary: str = ""
    project_objectives: str = ""
    target_beneficiaries: str = ""
    direct_beneficiaries: str = ""
    indirect_beneficiaries: str = ""
    project_activities: str = ""
    expected_outputs: str = ""
    staffing_plan: str = ""
    project_management_approach: str = ""
    expected_outcomes: str = ""
    quantitative_indicators: str = ""
    qualitative_indicators: str = ""
    baseline_conditions: str = ""
    baseline_data_collection: str = ""
    success_measurement: str = ""
    community_engagement: str = ""
    approvals_status: str = ""
    elders_involvement: str = ""
    knowledge_keepers_involvement: str = ""
    youth_involvement: str = ""
    data_governance: str = ""
    cultural_safety: str = ""
    evidence_note: str = ""
    why_now: str = ""
    total_project_cost: Optional[float] = None
    budget_personnel: str = ""
    budget_professional_services: str = ""
    budget_equipment_materials: str = ""
    budget_travel_logistics: str = ""
    budget_training: str = ""
    budget_evaluation: str = ""
    budget_admin: str = ""
    budget_contingency: str = ""
    budget_breakdown: str = ""
    budget_assumptions: str = ""
    other_funding: str = ""
    risks_and_mitigation: str = ""
    risk_likelihood: str = ""
    risk_impact: str = ""
    mitigation_plan: str = ""
    sustainability_plan: str = ""
    maintenance_requirements: str = ""
    ownership_model: str = ""
    future_funding_sources: str = ""
    scaling_plan: str = ""
    supporting_documents_text: str = ""
    requested_budget: Optional[float] = None
    budget_line_items: List[Dict[str, Any]] = Field(default_factory=list)
    budget_contingency_rate: Optional[float] = None
    budget_admin_rate: Optional[float] = None
    budget_participant_count: Optional[float] = None
    indicators_before: Optional[Dict[str, Any]] = None
    indicators_after: Optional[Dict[str, Any]] = None
    scenario: Optional[Dict[str, Any]] = None


class SectionSpec(BaseModel):
    key: str
    title: str
    guidance: str = ""
    word_limit: Optional[int] = None
    prompt_items: List[Dict[str, Any]] = Field(default_factory=list)


class RequirementsBody(BaseModel):
    grant_name: str = ""
    program_name: Optional[str] = None
    name: Optional[str] = None
    sections: List[SectionSpec] = []
    eligibility: List[str] = []
    word_limits: Dict[str, int] = {}
    must_include: List[str] = []
    raw_text: str = ""
    required_sections: List[str] = []


class GenerateDraftRequest(BaseModel):
    profile: CommunityProfile
    requirements: Dict[str, Any]  # full requirements as returned by parse
    requested_budget: float = Field(..., ge=0)


class EnhanceRequest(BaseModel):
    draft: Dict[str, Any]
    requirements: Dict[str, Any]
    profile: Dict[str, Any]
    use_case: Optional[str] = None


class ValidateRequest(BaseModel):
    draft: Dict[str, Any]
    requirements: Dict[str, Any]


class RewriteSectionRequest(BaseModel):
    section_key: str
    section_title: str = ""
    current_text: str = ""
    instruction: str = Field(..., min_length=1)
    requirements: Dict[str, Any]
    profile: Dict[str, Any]
    use_case: Optional[str] = None


class ExportSection(BaseModel):
    key: str = ""
    title: str
    body: str


class ExportDraftPdfRequest(BaseModel):
    grant_name: str = ""
    community_name: str = ""
    region: str = ""
    local_priority: str = ""
    requested_budget: Optional[int] = None
    sections: List[ExportSection] = []


class ExportDraftDocxRequest(ExportDraftPdfRequest):
    pass


class WorkspaceUser(BaseModel):
    id: str
    email: str
    name: str


class ProposalRecordRequest(BaseModel):
    title: Optional[str] = None
    community_name: Optional[str] = None
    grant_name: Optional[str] = None
    status: Optional[str] = None
    current_step: Optional[int] = None
    requirements: Optional[Dict[str, Any]] = None
    profile: Optional[Dict[str, Any]] = None
    draft: Optional[Dict[str, Any]] = None
    enhanced: Optional[Dict[str, Any]] = None
    structured_answers: Optional[Dict[str, Any]] = None
    prompt_coverage: Optional[Dict[str, Any]] = None
    validation: Optional[Dict[str, Any]] = None
    final_sections: Optional[List[Dict[str, Any]]] = None
    last_exported_at: Optional[str] = None
    community_profile_id: Optional[str] = None
    community_profile_snapshot: Optional[Dict[str, Any]] = None
    application_details: Optional[Dict[str, Any]] = None


class ProposalRecord(ProposalRecordRequest):
    id: str
    user_id: str
    title: str
    community_name: str = ""
    grant_name: str = ""
    status: str = "draft"
    current_step: int = 1
    created_at: str
    updated_at: str


class ProposalListResponse(BaseModel):
    proposals: List[ProposalRecord]


class CommunityProfileRequest(BaseModel):
    profile: Dict[str, Any] = Field(default_factory=dict)


class CommunityProfileRecord(CommunityProfileRequest):
    id: str
    user_id: str
    created_at: str
    updated_at: str


class FeedbackReportRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=200)
    source_filename: str = Field(default="", max_length=255)
    source_proposal_id: Optional[str] = None
    parent_report_id: Optional[str] = None
    overall_score: Optional[float] = Field(default=None, ge=0, le=100)
    priority_issue_count: int = Field(default=0, ge=0)
    category_scores: Dict[str, float] = Field(default_factory=dict)
    report: Dict[str, Any] = Field(default_factory=dict)
    extracted_sections: List[Dict[str, Any]] = Field(default_factory=list)
    grant_context: Optional[Dict[str, Any]] = None


class FeedbackReportRecord(FeedbackReportRequest):
    id: str
    user_id: str
    status: str
    rubric_version: str
    created_at: str
    analyzed_at: str


class FeedbackReportListResponse(BaseModel):
    reports: List[FeedbackReportRecord]


def _feedback_report_payload(
    analysis: ProposalAnalysisResponse,
    *,
    title: str,
    source_filename: str = "",
    source_proposal_id: Optional[str] = None,
    grant_context: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    analysis_payload = analysis.model_dump(mode="json")
    priority_issue_count = sum(
        1
        for category in analysis.categories
        for metric in category.metrics
        for issue in metric.issues
        if issue.severity in {"warning", "critical"}
    )
    return {
        "title": title,
        "source_filename": source_filename,
        "source_proposal_id": source_proposal_id,
        "status": "complete",
        "rubric_version": analysis.rubric_version,
        "overall_score": analysis.overall_score,
        "priority_issue_count": priority_issue_count,
        "category_scores": {category.id: category.score for category in analysis.categories},
        "report": {"analysis": analysis_payload},
        "extracted_sections": [section.model_dump(mode="json") for section in analysis.sections],
        "grant_context": grant_context,
        "analyzed_at": analysis.analysis.last_analyzed_at.isoformat(),
    }


def current_user(authorization: Optional[str] = Header(default=None)) -> Dict[str, Any]:
    try:
        user_payload = user_from_authorization_header(authorization)
    except ValueError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=401, detail=f"Authentication failed: {exc}") from exc
    return get_or_create_user(
        user_id=user_payload["id"],
        email=user_payload.get("email") or "user@example.com",
        name=user_payload.get("name") or "User",
    )


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_\- ]+", "", (value or "").strip())
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned[:60] or "grant_proposal"


_PROMPT_EXPORT_RE = re.compile(r"^\s*((?:Q[\w.-]+)|(?:prompt_\d+)|(?:\d[\w.-]*)):\s*(.+?)\s*$", re.IGNORECASE)


def _clean_export_answer(text: str) -> str:
    cleaned = (text or "").strip()
    cleaned = re.sub(r"(?im)^\s*Confidence:\s*(?:high|medium|low)\s*$", "", cleaned)
    cleaned = re.sub(r"(?im)^\s*Needs review:\s*.*$", "", cleaned)
    cleaned = cleaned.replace("[No answer generated]", "Needs additional information.")
    cleaned = cleaned.replace("[Missing information needed]", "Needs additional information.")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip() or "Needs additional information."


def _format_rag_counts(counts: Dict[str, Any] | None) -> str:
    if not counts:
        return "none"
    return ",".join(f"{key}:{value}" for key, value in counts.items()) or "none"


def _export_blocks(section_text: str) -> List[Dict[str, str]]:
    """Convert internal prompt_id blocks into clean export blocks."""
    lines = (section_text or "").replace("\r\n", "\n").split("\n")
    blocks: List[Dict[str, str]] = []
    current_label: Optional[str] = None
    current_answer: List[str] = []
    saw_prompt_blocks = False

    def flush() -> None:
        nonlocal current_label, current_answer
        if current_label is None:
            return
        blocks.append({
            "label": current_label,
            "body": _clean_export_answer("\n".join(current_answer)),
        })
        current_label = None
        current_answer = []

    for line in lines:
        match = _PROMPT_EXPORT_RE.match(line)
        if match:
            saw_prompt_blocks = True
            flush()
            prompt_id = match.group(1).strip()
            prompt_text = match.group(2).strip().rstrip(".")
            current_label = f"{prompt_id}: {prompt_text}"
            current_answer = []
            continue
        if current_label is not None:
            current_answer.append(line)

    flush()

    if saw_prompt_blocks:
        return blocks

    paragraphs = [
        _clean_export_answer(paragraph)
        for paragraph in re.split(r"\n\s*\n", section_text or "")
        if paragraph.strip()
    ]
    return [{"body": paragraph} for paragraph in paragraphs] or [{"body": "No content provided."}]


def _render_pdf(body: ExportDraftPdfRequest) -> bytes:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"PDF export dependency missing: {e}. Install reportlab.",
        )

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER
    left = 72
    right = width - 72
    top = height - 72
    bottom = 72
    line_h = 14
    y = top

    def new_page() -> None:
        nonlocal y
        c.showPage()
        y = top

    def ensure_room(lines: int = 1) -> None:
        nonlocal y
        if y - (lines * line_h) < bottom:
            new_page()

    def draw_line(text: str, *, font: str = "Helvetica", size: int = 11, extra_gap: float = 0.0) -> None:
        nonlocal y
        ensure_room(1)
        c.setFont(font, size)
        c.drawString(left, y, text)
        y -= line_h + extra_gap

    def wrap_text(text: str, size: int = 11) -> List[str]:
        from reportlab.pdfbase.pdfmetrics import stringWidth

        max_w = right - left
        words = (text or "").split()
        if not words:
            return [""]

        lines: List[str] = []
        cur = words[0]
        for w in words[1:]:
            candidate = f"{cur} {w}"
            if stringWidth(candidate, "Helvetica", size) <= max_w:
                cur = candidate
            else:
                lines.append(cur)
                cur = w
        lines.append(cur)
        return lines

    def draw_paragraph(text: str, *, size: int = 11, gap: float = 4.0) -> None:
        nonlocal y
        chunks = [p.strip() for p in (text or "").replace("\r\n", "\n").split("\n")]
        for chunk in chunks:
            lines = wrap_text(chunk, size=size)
            for ln in lines:
                ensure_room(1)
                c.setFont("Helvetica", size)
                c.drawString(left, y, ln)
                y -= line_h
            y -= gap

    def draw_subheading(text: str) -> None:
        nonlocal y
        y -= 2
        for ln in wrap_text(text, size=11):
            ensure_room(1)
            c.setFont("Helvetica-Bold", 11)
            c.drawString(left, y, ln)
            y -= line_h
        y -= 1

    # Cover header
    draw_line("Grant Proposal", font="Helvetica-Bold", size=22, extra_gap=8)
    if body.grant_name:
        draw_line(body.grant_name, font="Helvetica-Bold", size=14, extra_gap=6)
    draw_line(f"Community: {body.community_name or 'N/A'}", font="Helvetica", size=11)
    draw_line(f"Region: {body.region or 'N/A'}", font="Helvetica", size=11)
    draw_line(f"Local Priority: {body.local_priority or 'N/A'}", font="Helvetica", size=11)
    if body.requested_budget is not None:
        draw_line(
            f"Requested Funding: ${body.requested_budget:,.0f}",
            font="Helvetica",
            size=11,
        )
    draw_line(f"Generated: {date.today().isoformat()}", font="Helvetica", size=10, extra_gap=14)

    # Body sections
    for i, sec in enumerate(body.sections, start=1):
        if y < bottom + 120:
            new_page()
        title = (sec.title or f"Section {i}").strip()
        draw_line(f"{i}. {title}", font="Helvetica-Bold", size=13, extra_gap=3)
        for block in _export_blocks(sec.body or ""):
            label = block.get("label")
            if label:
                draw_subheading(label)
            draw_paragraph(block.get("body", ""), size=11, gap=6.0)

    c.save()
    buf.seek(0)
    return buf.read()


def _render_docx(body: ExportDraftDocxRequest) -> bytes:
    try:
        from docx import Document
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"DOCX export dependency missing: {e}. Install python-docx.",
        )

    document = Document()
    document.add_heading("Grant Proposal", level=0)
    if body.grant_name:
        document.add_paragraph(body.grant_name)
    document.add_paragraph(f"Community: {body.community_name or 'N/A'}")
    document.add_paragraph(f"Region: {body.region or 'N/A'}")
    document.add_paragraph(f"Local Priority: {body.local_priority or 'N/A'}")
    if body.requested_budget is not None:
        document.add_paragraph(f"Requested Funding: ${body.requested_budget:,.0f}")
    document.add_paragraph(f"Generated: {date.today().isoformat()}")

    for i, sec in enumerate(body.sections, start=1):
        title = (sec.title or f"Section {i}").strip()
        document.add_heading(f"{i}. {title}", level=1)
        for block in _export_blocks(sec.body or ""):
            label = block.get("label")
            if label:
                label_paragraph = document.add_paragraph()
                label_run = label_paragraph.add_run(label)
                label_run.bold = True
            document.add_paragraph(block.get("body", "No content provided."))

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf.read()


# ---------- Endpoints ----------
@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/api/me", response_model=WorkspaceUser)
def get_me(user: Dict[str, Any] = Depends(current_user)):
    return user


@app.get("/api/community-profile", response_model=CommunityProfileRecord)
def get_saved_community_profile(user: Dict[str, Any] = Depends(current_user)):
    profile = store_get_community_profile(user["id"])
    if not profile:
        raise HTTPException(status_code=404, detail="Community profile not found.")
    return profile


@app.put("/api/community-profile", response_model=CommunityProfileRecord)
def save_community_profile(body: CommunityProfileRequest, user: Dict[str, Any] = Depends(current_user)):
    return store_upsert_community_profile(user["id"], body.profile)


@app.get("/api/proposals", response_model=ProposalListResponse)
def list_saved_proposals(user: Dict[str, Any] = Depends(current_user)):
    return {"proposals": store_list_proposals(user["id"])}


@app.post("/api/proposals", response_model=ProposalRecord)
def create_saved_proposal(body: ProposalRecordRequest, user: Dict[str, Any] = Depends(current_user)):
    payload = body.model_dump(exclude_unset=True)
    return store_create_proposal(user["id"], payload)


@app.get("/api/proposals/{proposal_id}", response_model=ProposalRecord)
def get_saved_proposal(proposal_id: str, user: Dict[str, Any] = Depends(current_user)):
    proposal = store_get_proposal(user["id"], proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found.")
    return proposal


@app.patch("/api/proposals/{proposal_id}", response_model=ProposalRecord)
def update_saved_proposal(proposal_id: str, body: ProposalRecordRequest, user: Dict[str, Any] = Depends(current_user)):
    payload = body.model_dump(exclude_unset=True)
    proposal = store_update_proposal(user["id"], proposal_id, payload)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found.")
    return proposal


@app.post("/api/proposals/{proposal_id}/exported", response_model=ProposalRecord)
def mark_saved_proposal_exported(proposal_id: str, user: Dict[str, Any] = Depends(current_user)):
    proposal = mark_proposal_exported(user["id"], proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found.")
    return proposal


@app.post("/api/proposals/{proposal_id}/duplicate", response_model=ProposalRecord)
def duplicate_saved_proposal(proposal_id: str, user: Dict[str, Any] = Depends(current_user)):
    proposal = store_duplicate_proposal(user["id"], proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found.")
    return proposal


@app.delete("/api/proposals/{proposal_id}")
def delete_saved_proposal(proposal_id: str, user: Dict[str, Any] = Depends(current_user)):
    deleted = store_delete_proposal(user["id"], proposal_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Proposal not found.")
    return {"deleted": True}


@app.get("/api/proposal-feedback-reports", response_model=FeedbackReportListResponse)
def list_saved_feedback_reports(user: Dict[str, Any] = Depends(current_user)):
    return {"reports": store_list_feedback_reports(user["id"])}


@app.post("/api/proposal-feedback-reports", response_model=FeedbackReportRecord)
def create_saved_feedback_report(body: FeedbackReportRequest, user: Dict[str, Any] = Depends(current_user)):
    if body.source_proposal_id and not store_get_proposal(user["id"], body.source_proposal_id):
        raise HTTPException(status_code=404, detail="Source proposal not found.")
    if any(score < 0 or score > 100 for score in body.category_scores.values()):
        raise HTTPException(status_code=422, detail="Category scores must be between 0 and 100.")
    parent_report = None
    if body.parent_report_id:
        parent_report = store_get_feedback_report(user["id"], body.parent_report_id)
        if not parent_report:
            raise HTTPException(status_code=404, detail="Parent feedback report not found.")
        parent_source_id = parent_report.get("source_proposal_id")
        if body.source_proposal_id and parent_source_id and body.source_proposal_id != parent_source_id:
            raise HTTPException(status_code=422, detail="Parent report belongs to a different proposal.")
    payload = body.model_dump(exclude_unset=True)
    if parent_report and not payload.get("source_proposal_id"):
        payload["source_proposal_id"] = parent_report.get("source_proposal_id")
    payload.update({"status": "complete", "rubric_version": RUBRIC_VERSION})
    return store_create_feedback_report(user["id"], payload)


@app.post("/api/proposal-feedback-reports/analyze-upload", response_model=FeedbackReportRecord)
async def analyze_uploaded_proposal_for_feedback(
    proposal_file: UploadFile = File(...),
    grant_file: Optional[UploadFile] = File(default=None),
    user: Dict[str, Any] = Depends(current_user),
):
    if not proposal_file.filename:
        raise HTTPException(status_code=400, detail="A PDF or DOCX proposal draft is required.")
    proposal_name = Path(proposal_file.filename).name[:255]
    if Path(proposal_name).suffix.lower() not in {".pdf", ".docx"}:
        raise HTTPException(status_code=400, detail="Proposal drafts must be PDF or DOCX files.")
    proposal_content = await proposal_file.read()
    if len(proposal_content) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Proposal file exceeds the 15 MB upload limit.")

    grant_context = None
    if grant_file is not None and grant_file.filename:
        grant_name = Path(grant_file.filename).name[:255]
        if Path(grant_name).suffix.lower() not in {".pdf", ".docx", ".txt"}:
            raise HTTPException(status_code=400, detail="Grant guidelines must be a PDF, DOCX, or TXT file.")
        grant_content = await grant_file.read()
        if len(grant_content) > 15 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="Grant-guideline file exceeds the 15 MB upload limit.")
        try:
            from backend.app.parsers.grant_parsers import parse_grant_upload_to_requirements

            class ReviewUpload:
                def __init__(self, content: bytes, filename: str):
                    self._stream = BytesIO(content)
                    self.name = filename

                def read(self, size: int = -1):
                    return self._stream.read(size)

                def seek(self, offset: int, whence: int = 0):
                    return self._stream.seek(offset, whence)

                def tell(self):
                    return self._stream.tell()

                def getvalue(self):
                    return self._stream.getvalue()

            grant_requirements, grant_text = parse_grant_upload_to_requirements(ReviewUpload(grant_content, grant_name))
            grant_text = grant_text.strip()
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Could not extract the grant guidelines: {exc}") from exc
        if not grant_text:
            raise HTTPException(status_code=400, detail="No readable text was found in the grant guidelines.")
        grant_context = {
            "file_name": grant_name,
            "text": grant_text[:50000],
            "requirements": grant_requirements or {},
        }

    try:
        analysis = get_proposal_analysis_service().analyze_upload(
            proposal_name,
            proposal_content,
            persist_to_legacy_cache=False,
            grant_context_available=grant_context is not None,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    payload = _feedback_report_payload(
        analysis,
        title=f"{Path(proposal_name).stem[:180]} Review",
        source_filename=proposal_name,
        grant_context=grant_context,
    )
    return store_create_feedback_report(user["id"], payload)


@app.post("/api/proposals/{proposal_id}/feedback-report", response_model=FeedbackReportRecord)
def analyze_saved_proposal_for_feedback(proposal_id: str, user: Dict[str, Any] = Depends(current_user)):
    proposal = store_get_proposal(user["id"], proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found.")
    source_sections = proposal.get("final_sections") or []
    if not source_sections:
        raise HTTPException(status_code=422, detail="Generate and review the proposal before requesting feedback.")
    sections = [
        ProposalSection(
            key=str(section.get("key") or f"section_{index + 1}"),
            title=str(section.get("title") or f"Section {index + 1}"),
            body=str(section.get("body") or ""),
            order=index,
            word_limit=section.get("word_limit"),
        )
        for index, section in enumerate(source_sections)
    ]
    title = proposal.get("title") or proposal.get("grant_name") or "Proposal"
    analysis = get_proposal_analysis_service().analyze_sections_snapshot(
        file_name=title,
        sections=sections,
        grant_context_available=bool(proposal.get("requirements")),
    )
    requirements = proposal.get("requirements") or {}
    grant_context = {
        "file_name": requirements.get("grant_name") or proposal.get("grant_name") or "Saved grant requirements",
        "text": str(requirements.get("raw_text") or "")[:50000],
        "requirements": requirements,
    } if requirements else None
    payload = _feedback_report_payload(
        analysis,
        title=f"{title} Review",
        source_proposal_id=proposal_id,
        grant_context=grant_context,
    )
    return store_create_feedback_report(user["id"], payload)


@app.get("/api/proposal-feedback-reports/{report_id}", response_model=FeedbackReportRecord)
def get_saved_feedback_report(report_id: str, user: Dict[str, Any] = Depends(current_user)):
    report = store_get_feedback_report(user["id"], report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Feedback report not found.")
    return report


@app.delete("/api/proposal-feedback-reports/{report_id}")
def delete_saved_feedback_report(report_id: str, user: Dict[str, Any] = Depends(current_user)):
    deleted = store_delete_feedback_report(user["id"], report_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Feedback report not found.")
    return {"deleted": True}


def get_compliance_service() -> ComplianceEvaluationService:
    service = getattr(app.state, "compliance_service", None)
    if service is None:
        service = build_default_service()
        app.state.compliance_service = service
    return service


def get_proposal_analysis_service() -> ProposalAnalysisService:
    service = getattr(app.state, "proposal_analysis_service", None)
    if service is None:
        service = ProposalAnalysisService(get_compliance_service())
        app.state.proposal_analysis_service = service
    return service


@app.post("/evaluate/compliance", response_model=ComplianceEvaluationResponse)
def evaluate_compliance(body: ComplianceEvaluationRequest):
    return get_compliance_service().evaluate_section(body)


@app.post("/evaluate/proposal", response_model=ProposalAnalysisResponse)
async def evaluate_proposal(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="A PDF or DOCX file is required.")
    lower = file.filename.lower()
    if not (lower.endswith(".pdf") or lower.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
    try:
        content = await file.read()
        analysis = get_proposal_analysis_service().analyze_upload(file.filename, content)
        return analysis
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/evaluate/proposal/reanalyze", response_model=ProposalAnalysisResponse)
def reanalyze_proposal(body: ProposalReanalyzeRequest):
    try:
        return get_proposal_analysis_service().reanalyze_sections(
            proposal_id=body.proposal_id,
            sections=body.sections,
        )
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail="Proposal analysis not found.") from exc


@app.get("/evaluate/proposal/{proposal_id}", response_model=ProposalAnalysisResponse)
def get_proposal_analysis(proposal_id: str):
    try:
        return get_proposal_analysis_service().load_analysis(proposal_id)
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail="Proposal analysis not found.") from exc


@app.post("/evaluate/proposal/section-rewrite", response_model=ProposalSectionRewriteResponse)
def rewrite_proposal_section(body: ProposalSectionRewriteRequest):
    try:
        return get_proposal_analysis_service().rewrite_section(
            proposal_id=body.proposal_id,
            section_key=body.section_key,
            instruction=body.instruction,
            rewrite_scope=body.rewrite_scope,
            target_text=body.target_text,
            metric_id=body.metric_id,
            issue_id=body.issue_id,
            issue_message=body.issue_message,
            issue_recommendation=body.issue_recommendation,
        )
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail="Proposal analysis not found.") from exc
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="Section not found.") from exc


@app.post("/evaluate/proposal/chat", response_model=ProposalChatResponse)
def proposal_chat(body: ProposalChatRequest):
    try:
        return get_proposal_analysis_service().chat(
            proposal_id=body.proposal_id,
            message=body.message,
            section_key=body.section_key,
            metric_id=body.metric_id,
        )
    except FileNotFoundError as exc:
        raise HTTPException(status_code=404, detail="Proposal analysis not found.") from exc


@app.post("/api/parse-grant")
async def parse_grant(file: UploadFile = File(...)):
    """Parse uploaded grant document (PDF/DOCX/TXT) into structured requirements."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename")
    name = file.filename.lower()
    if not (name.endswith(".txt") or name.endswith(".pdf") or name.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only .txt, .pdf, .docx supported")
    try:
        content = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {e}")
    # Parser expects a file-like with .name and .read() / .getvalue()
    class FileLike:
        def __init__(self, data: bytes, filename: str):
            self._io = BytesIO(data)
            self.name = filename
        def read(self, n=-1):
            return self._io.read(n)
        def getvalue(self):
            return self._io.getvalue()
    file_like = FileLike(content, file.filename)
    try:
        from backend.app.parsers.grant_parsers import parse_grant_upload_to_requirements
        requirements, raw_text = parse_grant_upload_to_requirements(file_like)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if requirements is None:
        raise HTTPException(status_code=422, detail="Could not parse document")
    # Convert sections to serializable dicts
    req_dict = {
        "grant_name": requirements.get("grant_name"),
        "sections": [
            {
                "key": s.get("key"),
                "title": s.get("title"),
                "guidance": s.get("guidance", ""),
                "word_limit": s.get("word_limit"),
                "prompt_items": s.get("prompt_items", []),
                "section_purpose": s.get("section_purpose"),
                "parser_diagnostics": s.get("parser_diagnostics", []),
            }
            for s in requirements.get("sections", [])
        ],
        "eligibility": requirements.get("eligibility", []),
        "word_limits": requirements.get("word_limits", {}),
        "must_include": requirements.get("must_include", []),
        "raw_text": requirements.get("raw_text", raw_text),
        "required_sections": requirements.get("required_sections", []),
        "parser_meta": requirements.get("parser_meta", {}),
    }
    return {"requirements": req_dict, "raw_text": raw_text}


@app.post("/api/parse-supporting-document")
async def parse_supporting_document(file: UploadFile = File(...)):
    """Extract raw text from a supporting context document."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename")
    name = file.filename.lower()
    if not name.endswith((".txt", ".md", ".csv", ".json", ".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="Only .txt, .md, .csv, .json, .pdf, .docx supported")
    try:
        content = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {e}")

    class FileLike:
        def __init__(self, data: bytes, filename: str):
            self._io = BytesIO(data)
            self.name = filename
        def read(self, n=-1):
            return self._io.read(n)
        def getvalue(self):
            return self._io.getvalue()
        def seek(self, pos=0):
            return self._io.seek(pos)

    try:
        from backend.app.parsers.grant_parsers import extract_text_from_upload
        raw_text = extract_text_from_upload(FileLike(content, file.filename)).strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not raw_text:
        raise HTTPException(status_code=422, detail="Could not extract text from supporting document")
    return {
        "filename": file.filename,
        "raw_text": raw_text[:20000],
        "char_count": len(raw_text),
    }


@app.post("/api/generate-draft")
def generate_draft(body: GenerateDraftRequest):
    """Generate a baseline draft from community profile and requirements."""
    try:
        from backend.app.utils.grant_utils import generate_proposal_from_requirements
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"Backend import error: {e}")
    profile = body.profile.model_dump()
    requirements = body.requirements
    budget = body.requested_budget
    profile["requested_budget"] = budget
    from backend.app.utils.budget_calculations import calculate_budget_outputs
    profile["verified_budget_calculations"] = calculate_budget_outputs(profile)
    draft = generate_proposal_from_requirements(profile=profile, requirements=requirements, requested_budget=budget)
    return draft


@app.post("/api/enhance")
def enhance(body: EnhanceRequest):
    """Enhance draft sections using RAG + LLM."""
    try:
        from backend.app.llm.llm_utils import enhance_sections_with_metadata
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"Backend import error: {e}")
    section_count = len((body.draft or {}).get("sections", []) or [])
    prompt_count = sum(
        len(section.get("prompt_items", []) or [])
        for section in (body.draft or {}).get("sections", []) or []
        if isinstance(section, dict)
    )
    logger.info(
        "Enhance request received sections=%s prompt_items=%s grant=%s model=%s openai_key_configured=%s",
        section_count,
        prompt_count,
        (body.requirements or {}).get("grant_name"),
        os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini"),
        bool(os.getenv("OPENAI_API_KEY")),
    )
    result = enhance_sections_with_metadata(
        draft=body.draft,
        requirements=body.requirements,
        profile=body.profile,
        use_case=body.use_case,
    )
    meta = result.get("meta", {}) if isinstance(result, dict) else {}
    logger.info(
        "Enhance request completed sections=%s enhanced_sections=%s fallback_used=%s fallback_reason=%s rag_collection=%s rag_vector_used=%s rag_vector_attempted=%s rag_vector_counts=%s rag_query_chars=%s rag_chars=%s rag_fallback_used=%s rag_fallback_reason=%s rag_error=%s structured_answers=%s",
        section_count,
        meta.get("enhanced_section_count"),
        meta.get("fallback_used"),
        meta.get("fallback_reason") or "none",
        meta.get("rag_collection") or "none",
        meta.get("rag_vector_collection_used") or "none",
        ",".join(meta.get("rag_vector_attempted_collections") or []) or "none",
        _format_rag_counts(meta.get("rag_vector_collection_counts")),
        meta.get("rag_query_chars"),
        meta.get("rag_context_chars"),
        meta.get("rag_fallback_used"),
        meta.get("rag_fallback_reason") or "none",
        meta.get("rag_error") or "none",
        meta.get("structured_answer_count"),
    )
    return result


@app.post("/api/validate")
def validate(body: ValidateRequest):
    """Validate draft against requirements."""
    try:
        from backend.app.utils.validation_utils import validate_proposal_against_requirements
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"Backend import error: {e}")
    result = validate_proposal_against_requirements(draft=body.draft, requirements=body.requirements)
    return result


@app.post("/api/rewrite-section")
def rewrite_section(body: RewriteSectionRequest):
    """Rewrite one draft section with user instruction and return source references."""
    try:
        from backend.app.llm.llm_utils import rewrite_section_with_instruction
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"Backend import error: {e}")

    try:
        out = rewrite_section_with_instruction(
            section_key=body.section_key,
            section_title=body.section_title,
            current_text=body.current_text,
            instruction=body.instruction,
            requirements=body.requirements,
            profile=body.profile,
            use_case=body.use_case,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    return out


@app.post("/api/export-draft-pdf")
def export_draft_pdf(body: ExportDraftPdfRequest):
    """Export the final draft as a professionally formatted PDF."""
    if not body.sections:
        raise HTTPException(status_code=400, detail="At least one section is required for export.")

    pdf_bytes = _render_pdf(body)
    base = _safe_filename(body.community_name or body.grant_name or "grant_proposal")
    filename = f"{base}_proposal.pdf"
    return StreamingResponse(
        BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/export-draft-docx")
def export_draft_docx(body: ExportDraftDocxRequest):
    """Export the final draft as a DOCX document."""
    if not body.sections:
        raise HTTPException(status_code=400, detail="At least one section is required for export.")

    docx_bytes = _render_docx(body)
    base = _safe_filename(body.community_name or body.grant_name or "grant_proposal")
    filename = f"{base}_proposal.docx"
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
